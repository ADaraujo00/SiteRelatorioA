import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from PIL import Image
import docx  # Importamos a biblioteca docx explicitamente

OUTPUT_FILENAME = 'relatorio_completo.docx'
IMAGE_SIZE_CM = 4
IMAGE_SIZE_INCHES = IMAGE_SIZE_CM / 2.54
# Salve a imagem com este nome na mesma pasta do script
HEADER_IMAGE_PATH = 'electrolux_header.png'

st.title("Gerador de Relatório Completo")

report_number = st.text_input("**Número do Relatório:**", "TR00001")

st.subheader("Informações Gerais:")

col1, col2, col3, col4 = st.columns(4)
product_general = col1.text_input("**Product:**")
project = col2.text_input("**Project:**")
lot = col3.text_input("**Lot:**")
released = col4.text_input("**Released:**")

col5, col6, col7, col8 = st.columns(4)
requested_by = col5.text_input("**Requested by:**")
performed_by = col6.text_input("**Performed by:**")
reviewed_by = col7.text_input("**Reviewed by:**")
approved_by = col8.text_input("**Approved by:**")

st.subheader("3. SAMPLES DESCRIPTION:")

image_labels = [
    "Foto base frente", "Foto base costas", "Foto base cima", "Foto base baixo",
    "Produto inteiro frente com jarra blender", "Produto inteiro lado com jarra blender",
    "Produto inteiro costas com jarra blender", "Produto inteiro lado com jarra blender",
    "Jarra lado", "Jarra frente", "Jarra cima", "Jarra embaixo"
]

image_uploaders = []
cols_grid = [st.columns(4) for _ in range(3)]
for i in range(3):
    for j in range(4):
        index = i * 4 + j
        uploader = cols_grid[i][j].file_uploader(f"{image_labels[index]}:", type=["jpg", "jpeg", "png"], key=f"image_{index}")
        image_uploaders.append(uploader)

st.subheader("Detalhes das Amostras:")
sample_details_input = {}
questions = ["Product", "Accessories", "Model", "Voltage", "Dimensions", "Supplier", "Quantity", "Volume"]

for question in questions:
    col_q, col_a = st.columns([1, 2])
    col_q.markdown(f"**{question}:**")
    sample_details_input[question] = col_a.text_input("", label_visibility="collapsed", key=question)

st.subheader("Análise da Peneira:")

col_labels_sieve = ["", "Minimum", "Medium", "Maximum"]
row_labels_sieve_input = ["Carrots mass (g)", "Water mass (g)", "Mass retained on sieve 2 mm (g)", "Mass retained on sieve 4 mm (g)"]
sieve_input_values = {}
performance_2mm = {}
performance_4mm = {}

cols_sieve_header = st.columns(len(col_labels_sieve))
for i, label in enumerate(col_labels_sieve):
    cols_sieve_header[i].markdown(f"**{label}**")

# Inputs do usuário para a análise da peneira
for i, row_label in enumerate(row_labels_sieve_input):
    cols_sieve_input = st.columns(len(col_labels_sieve))
    cols_sieve_input[0].markdown(f"**{row_label}**")
    for j in range(1, len(col_labels_sieve)):
        key = f"{row_label.replace(' ', '_')}_{col_labels_sieve[j].lower()}"
        sieve_input_values[key] = cols_sieve_input[j].text_input("", key=key)

# Cálculo da Performance da Peneira
for j in range(1, len(col_labels_sieve)):
    min_med_max = col_labels_sieve[j].lower()
    carrots_mass = float(sieve_input_values.get(f"Carrots_mass_(g)_{min_med_max}", 0) or 0)
    retained_2mm = float(sieve_input_values.get(f"Mass_retained_on_sieve_2_mm_(g)_{min_med_max}", 0) or 0)
    retained_4mm = float(sieve_input_values.get(f"Mass_retained_on_sieve_4_mm_(g)_{min_med_max}", 0) or 0)

    # Cálculo da Performance 2 mm
    if carrots_mass > 0:
        mr_2mm = retained_2mm + retained_4mm
        performance_2mm[min_med_max] = 1 - (mr_2mm / carrots_mass) if carrots_mass > 0 else 0
        # Cálculo da Performance 4 mm (mantendo o cálculo anterior)
        performance_4mm[min_med_max] = 1 - (retained_4mm / carrots_mass) if carrots_mass > 0 else 0
    else:
        performance_2mm[min_med_max] = 0
        performance_4mm[min_med_max] = 0

# Exibição das Performances da Peneira
st.markdown("---")
cols_performance_2mm = st.columns(len(col_labels_sieve))
cols_performance_2mm[0].markdown("**Performance 2 mm**")
for i in range(1, len(col_labels_sieve)):
    value = performance_2mm.get(col_labels_sieve[i].lower(), 0)
    cols_performance_2mm[i].markdown(f"{value:.2%}")

cols_performance_4mm = st.columns(len(col_labels_sieve))
cols_performance_4mm[0].markdown("**Performance 4 mm**")
for i in range(1, len(col_labels_sieve)):
    value = performance_4mm.get(col_labels_sieve[i].lower(), 0)
    cols_performance_4mm[i].markdown(f"{value:.2%}")

st.subheader("Fotos da Peneira:")

col_photo1, col_label1 = st.columns([1, 3])
photo_2mm = col_photo1.file_uploader("Foto Peneira 2 mm:", type=["jpg", "jpeg", "png"], key="uploader_photo_2mm")
col_label1.markdown("**Picture 2 mm**")

col_photo2, col_label2 = st.columns([1, 3])
photo_4mm = col_photo2.file_uploader("Foto Peneira 4 mm:", type=["jpg", "jpeg", "png"], key="uploader_photo_4mm")
col_label2.markdown("**Picture 4 mm**")

def generate_word_report(report_num, general_data, images_data, sample_data, sieve_data, perf_2mm, perf_4mm, sieve_photos):
    doc = Document()

    # Definir as margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.94)
        section.bottom_margin = Cm(1.62)
        section.left_margin = Cm(1.02)
        section.right_margin = Cm(1.02)

    # Definir a fonte Arial tamanho 10 como padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Adicionar a imagem de cabeçalho
    try:
        doc.add_picture(HEADER_IMAGE_PATH, width=Inches(19.50 / 2.54))
        header_paragraph = doc.paragraphs[-1]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"Erro ao adicionar imagem de cabeçalho: {e}")
        paragraph = doc.paragraphs[-1]
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = Pt(10)

    # Adicionar o número do relatório
    paragraph = doc.add_paragraph(f"Relatório No: {report_num}")
    paragraph.style.font.name = 'Arial'
    paragraph.style.font.size = Pt(10)
    doc.add_paragraph()  # Espaço

    def remove_table_borders(table):
        tblPr = table._element.xpath('./w:tblPr')[0]
        borders = docx.oxml.shared.OxmlElement('w:tblBorders')
        for tag in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = docx.oxml.shared.OxmlElement(f'w:{tag}')
            border.set(docx.oxml.ns.qn('w:val'), 'nil')
            borders.append(border)
        tblPr.append(borders)

    # Informações Gerais em tabela com respostas abaixo
    table_row1 = doc.add_table(rows=2, cols=4)
    for row in table_row1.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(10)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    if "Product:" in cell.text or "Project:" in cell.text or "Lot:" in cell.text or "Released:" in cell.text:
                        run.font.bold = True
    cells_row1 = table_row1.rows[0].cells
    cells_row1[0].text = "Product:"
    cells_row1[1].text = "Project:"
    cells_row1[2].text = "Lot:"
    cells_row1[3].text = "Released:"
    cells_row2 = table_row1.rows[1].cells
    cells_row2[0].text = general_data['Product']
    cells_row2[1].text = general_data['Project']
    cells_row2[2].text = general_data['Lot']
    cells_row2[3].text = general_data['Released']
    remove_table_borders(table_row1)

    table_row2 = doc.add_table(rows=2, cols=4)
    for row in table_row2.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(10)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    if "Requested by:" in cell.text or "Performed by:" in cell.text or "Reviewed by:" in cell.text or "Approved by:" in cell.text:
                        run.font.bold = True
    cells_row3 = table_row2.rows[0].cells
    cells_row3[0].text = "Requested by:"
    cells_row3[1].text = "Performed by:"
    cells_row3[2].text = "Reviewed by:"
    cells_row3[3].text = "Approved by:"
    cells_row4 = table_row2.rows[1].cells
    cells_row4[0].text = general_data['Requested by']
    cells_row4[1].text = general_data['Performed by']
    cells_row4[2].text = general_data['Reviewed by']
    cells_row4[3].text = general_data['Approved by']
    remove_table_borders(table_row2)

    doc.add_paragraph()  # Espaço antes da próxima seção

    doc.add_paragraph("5. SAMPLES DESCRIPTION:")
    paragraph = doc.paragraphs[-1]
    paragraph.style.font.name = 'Arial'
    paragraph.style.font.size = Pt(10)
    paragraph.runs[0].font.bold = True

    # Adicionar as imagens em tabela 3x4 com tamanho fixo
    image_table = doc.add_table(rows=3, cols=4)
    images = [img for img in images_data if img is not None]
    if images:  # Verifica se há alguma imagem carregada
        for i in range(3):
            for j in range(4):
                index = i * 4 + j
                cell = image_table.cell(i, j)
                if index < len(images):
                    try:
                        image_stream = BytesIO(images[index])
                        cell.paragraphs[0].add_run().add_picture(
                            image_stream,
                            width=Inches(IMAGE_SIZE_INCHES),
                            height=Inches(IMAGE_SIZE_INCHES)
                        )
                    except Exception as e:
                        cell.text = f"Erro ao adicionar imagem {index + 1}: {e}"
                        for paragraph in cell.paragraphs:
                            paragraph.style.font.name = 'Arial'
                            paragraph.style.font.size = Pt(10)
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                if index < len(image_labels):
                    cell.add_paragraph(image_labels[index]).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph = cell.paragraphs[-1]
                    paragraph.style.font.name = 'Arial'
                    paragraph.style.font.size = Pt(8)
    else:
        # Se não houver imagens, preenche as células com os rótulos
        for i in range(3):
            for j in range(4):
                index = i * 4 + j
                cell = image_table.cell(i, j)
                if index < len(image_labels):
                    cell.add_paragraph(image_labels[index]).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph = cell.paragraphs[-1]
                    paragraph.style.font.name = 'Arial'
                    paragraph.style.font.size = Pt(8)

    doc.add_paragraph("\n6. DETALHES DAS AMOSTRAS:")
    paragraph = doc.paragraphs[-1]
    paragraph.style.font.name = 'Arial'
    paragraph.style.font.size = Pt(10)
    paragraph.runs[0].font.bold = True
    table_samples = doc.add_table(rows=len(sample_data), cols=2)
    for i, (key, value) in enumerate(sample_data.items()):
        cell_left = table_samples.cell(i, 0)
        cell_right = table_samples.cell(i, 1)
        for paragraph in cell_left.paragraphs:
            paragraph.style.font.name = 'Arial'
            paragraph.style.font.size = Pt(10)
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
        cell_left.text = f"{key}:"
        for paragraph in cell_right.paragraphs:
            paragraph.style.font.name = 'Arial'
            paragraph.style.font.size = Pt(10)
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
        cell_right.text = value

    # Adicionar a seção de análise da peneira
    doc.add_paragraph("\n7. ANÁLISE DA PENEIRA:")
    paragraph = doc.paragraphs[-1]
    paragraph.style.font.name = 'Arial'
    paragraph.style.font.size = Pt(10)
    paragraph.runs[0].font.bold = True
    sieve_table = doc.add_table(rows=len(row_labels_sieve_input) + 2, cols=len(col_labels_sieve))
    # Cabeçalho
    for i, label in enumerate(col_labels_sieve):
        cell = sieve_table.cell(0, i)
        cell.text = label
        for paragraph in cell.paragraphs:
            paragraph.style.font.name = 'Arial'
            paragraph.style.font.size = Pt(10)
            paragraph.runs[0].font.bold = True

    # Linhas de input
    for i, row_label in enumerate(row_labels_sieve_input):
        cell = sieve_table.cell(i + 1, 0)
        cell.text = row_label
        for paragraph in cell.paragraphs:
            paragraph.style.font.name = 'Arial'
            paragraph.style.font.size = Pt(10)
            paragraph.runs[0].font.bold = True
        for j in range(1, len(col_labels_sieve)):
            key = f"{row_label.replace(' ', '_')}_{col_labels_sieve[j].lower()}"
            value = sieve_data.get(key, "")
            cell = sieve_table.cell(i + 1, j)
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(10)

    # Linhas de performance
    perf_2mm_row_index = len(row_labels_sieve_input) + 1
    cell_perf_2mm_label = sieve_table.cell(perf_2mm_row_index, 0)
    cell_perf_2mm_label.text = "Performance 2 mm"
    for paragraph in cell_perf_2mm_label.paragraphs:
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = Pt(10)
        paragraph.runs[0].font.bold = True
    for j in range(1, len(col_labels_sieve)):
        min_med_max = col_labels_sieve[j].lower()
        value = perf_2mm.get(min_med_max, 0)
        cell_perf_2mm_value = sieve_table.cell(perf_2mm_row_index, j)
        cell_perf_2mm_value.text = f"{value:.2%}"
        for paragraph in cell_perf_2mm_value.paragraphs:
            paragraph.style.font.name = 'Arial'
            paragraph.style.font.size = Pt(10)

    perf_4mm_row_index = len(row_labels_sieve_input) + 1
    cell_perf_4mm_label = sieve_table.cell(perf_4mm_row_index, 0)
    cell_perf_4mm_label.text = "Performance 4 mm"
    for paragraph in cell_perf_4mm_label.paragraphs:
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = Pt(10)
        paragraph.runs[0].font.bold = True
    for j in range(1, len(col_labels_sieve)):
        min_med_max = col_labels_sieve[j].lower()
        value = perf_4mm.get(min_med_max, 0)
        cell_perf_4mm_value = sieve_table.cell(perf_4mm_row_index, j)
        cell_perf_4mm_value.text = f"{value:.2%}"
        for paragraph in cell_perf_4mm_value.paragraphs:
            paragraph.style.font.name = 'Arial'
            paragraph.style.font.size = Pt(10)

    remove_table_borders(sieve_table)
    doc.add_paragraph()

    # Adicionar fotos da peneira
    doc.add_paragraph("\n8. FOTOS DA PENEIRA:")
    paragraph = doc.paragraphs[-1]
    paragraph.style.font.name = 'Arial'
    paragraph.style.font.size = Pt(10)
    paragraph.runs[0].font.bold = True

    sieve_photo_table = doc.add_table(rows=2, cols=2)
    if sieve_photos.get("photo_2mm"):
        cell_2mm = sieve_photo_table.cell(0, 0)
        try:
            image_stream_2mm = BytesIO(sieve_photos["photo_2mm"])
            cell_2mm.paragraphs[0].add_run().add_picture(
                image_stream_2mm,
                width=Inches(3),
                height=Inches(3)
            )
            cell_2mm.add_paragraph("Peneira 2 mm").alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell_2mm.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(8)
        except Exception as e:
            cell_2mm.text = f"Erro ao adicionar foto 2mm: {e}"
            for paragraph in cell_2mm.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(10)

    if sieve_photos.get("photo_4mm"):
        cell_4mm = sieve_photo_table.cell(0, 1)
        try:
            image_stream_4mm = BytesIO(sieve_photos["photo_4mm"])
            cell_4mm.paragraphs[0].add_run().add_picture(
                image_stream_4mm,
                width=Inches(3),
                height=Inches(3)
            )
            cell_4mm.add_paragraph("Peneira 4 mm").alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell_4mm.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(8)
        except Exception as e:
            cell_4mm.text = f"Erro ao adicionar foto 4mm: {e}"
            for paragraph in cell_4mm.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(10)

    # Salvar o documento na memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


if st.button("Gerar Relatório"):
    uploaded_files = [uploader for uploader in image_uploaders if uploader is not None]
    general_info = {
        "Product": product_general,
        "Project": project,
        "Lot": lot,
        "Released": released,
        "Requested by": requested_by,
        "Performed by": performed_by,
        "Reviewed by": reviewed_by,
        "Approved by": approved_by
    }
    sieve_data_report = {}
    for key, value in sieve_input_values.items():
        sieve_data_report[key] = value

    sieve_photos_report = {
        "photo_2mm": photo_2mm.getvalue() if photo_2mm else None,
        "photo_4mm": photo_4mm.getvalue() if photo_4mm else None,
    }

    word_buffer = generate_word_report(
        report_number,
        general_info,
        [uf.getvalue() if uf else None for uf in uploaded_files],
        sample_details_input,
        sieve_data_report,
        performance_2mm,
        performance_4mm,
        sieve_photos_report
    )
    if word_buffer:
        st.download_button(
            label="Baixar Relatório Word",
            data=word_buffer.getvalue(),
            file_name=OUTPUT_FILENAME,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

st.markdown("""
---
**Instruções:**
1. Insira o **Número do Relatório**.
2. Preencha as **Informações Gerais**.
3. Carregue as imagens das amostras.
4. Preencha os **Detalhes das Amostras**.
5. Preencha a **Análise da Peneira**.
6. Carregue as fotos da peneira.
7. Clique em "Gerar Relatório" para baixar o documento Word.
""")
